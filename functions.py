import mysql.connector
from passlib.hash import bcrypt
import os
import uuid
import shutil
from datetime import datetime
import time
from html_content import *
from fastapi.responses import HTMLResponse
from fastapi import HTTPException, status
import boto3
import configparser
import requests
import pandas as pd
import docx
import random
from deep_translator import GoogleTranslator


class Connect:
    def config():
        global config
        config = configparser.ConfigParser()
        config.read("config.ini", encoding='utf-8')
        return config

    def initiate_db():
        config = Connect.config()
        db = mysql.connector.connect(
            host=config['MySQL']['host'],
            user=config['MySQL']['user'],
            passwd=config['MySQL']['passwd'],
            port=config['MySQL']['port']
        )
        database_cursor = db.cursor()
        database_cursor.execute("show databases")
        databases = []
        for i in database_cursor:
            databases.append(i[0])
        if "transcriptions" not in databases:
            start_cursor = db.cursor()
            start_cursor.execute("CREATE DATABASE transcriptions;")
            start_cursor.execute("CREATE TABLE `transcriptions`.`Users` (`id` INT NOT NULL AUTO_INCREMENT,`email_id` VARCHAR(50) NULL,`username` VARCHAR(50) NULL,`password` VARCHAR(150) NULL,PRIMARY KEY (`id`),UNIQUE INDEX `email_id_UNIQUE` (`email_id` ASC) VISIBLE,UNIQUE INDEX `username_UNIQUE` (`username` ASC) VISIBLE);")
            start_cursor.execute("CREATE TABLE `transcriptions`.`Records` (`id` INT NOT NULL AUTO_INCREMENT,`date_time` DATETIME NULL,`username` VARCHAR(50) NULL,`project_name` VARCHAR(200) NULL,`fname` VARCHAR(200) NULL,`doc_download_link` VARCHAR(500) NULL,`csv_download_link` VARCHAR(500) NULL,`transcription_time` VARCHAR(50) NULL,PRIMARY KEY (`id`));")
            start_cursor.execute(
                "CREATE TABLE `transcriptions`.`Projects` (`id` INT NOT NULL AUTO_INCREMENT,`project_name` VARCHAR(200) NULL,`username` VARCHAR(50) NULL,PRIMARY KEY (`id`),UNIQUE INDEX `project_name_UNIQUE` (`project_name` ASC) VISIBLE);")
            start_cursor.execute(
                "CREATE TABLE `transcriptions`.`Languages` (`id` INT NOT NULL AUTO_INCREMENT,`language` VARCHAR(200) NULL,`language_code` VARCHAR(5) NULL,PRIMARY KEY (`id`),UNIQUE INDEX `language_UNIQUE` (`language` ASC) VISIBLE);")
            db.commit()

    def db():
        global mydb
        config = Connect.config()
        mydb = mysql.connector.connect(
            host=config['MySQL']['host'],
            user=config['MySQL']['user'],
            passwd=config['MySQL']['passwd'],
            port=config['MySQL']['port'],
            database=config['MySQL']['database']
        )
        return mydb

    def s3():
        global s3
        config = Connect.config()
        s3 = boto3.client('s3',
                          aws_access_key_id=config['S3SETTINGS']['access_key_id'],
                          aws_secret_access_key=config['S3SETTINGS']['secret_key_id'],
                          region_name=config['S3SETTINGS']['region'])
        return s3


def authenticate(username, password):
    mydb = Connect.db()
    user_cursor = mydb.cursor()
    try:
        user_cursor.execute(
            "select password from Users where username = '"+username+"';")
        stored_pass = user_cursor.fetchall()[0][0]
        auth = bcrypt.verify(password, stored_pass)
    except IndexError:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED,
                            detail='Invalid username or password')
    return True


def get_english_projects():
    mydb = Connect.db()
    project_cursor = mydb.cursor()
    project_cursor.execute(
        "select project_name from Projects WHERE language ='English';")
    select_box_html = ""
    for i in project_cursor:
        select_box_html += "<option value='"+i[0]+"'>"+i[0]+"</option>"
    dir_name = str(uuid.uuid4())
    return select_box_html, dir_name


def get_other_projects():
    mydb = Connect.db()
    project_cursor = mydb.cursor()
    project_cursor.execute(
        "select project_name from Projects WHERE language ='Non English';")
    select_box_html = ""
    for i in project_cursor:
        select_box_html += "<option value='"+i[0]+"'>"+i[0]+"</option>"

    project_cursor.execute("select language from Languages;")
    language_box_html = ""
    for i in project_cursor:
        language_box_html += "<option value='"+i[0]+"'>"+i[0]+"</option>"

    dir_name = str(uuid.uuid4())
    return select_box_html, language_box_html, dir_name


def emotion_results(dict, emotion):
    try:
        score = dict[emotion]
        return score
    except:
        return None

# ------------------------------------------


def transcribe_english(file, token):
    headers = {'authorization': token}
    url_response = requests.post('https://api.assemblyai.com/v2/upload',
                                 headers=headers,
                                 data=file)

    url = url_response.json()["upload_url"]

    id_endpoint = "https://api.assemblyai.com/v2/transcript"

    json = {
        "audio_url": url,
        "speaker_labels": True,
        "auto_highlights": True,
        "iab_categories": True,
        "sentiment_analysis": True,
        "auto_chapters": True,
        "entity_detection": True
    }

    headers = {"authorization": token, "content-type": "application/json"}
    id_response = requests.post(id_endpoint, json=json, headers=headers)
    transcribe_id = id_response.json()['id']
    text_endpoint = "https://api.assemblyai.com/v2/transcript/"+transcribe_id
    headers = {"authorization": token, }
    result = requests.get(text_endpoint, headers=headers).json()

    while result.get("status") != 'completed':
        if result.get("status") == 'error':
            return "error"
        text_endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {"authorization": token}
        result = requests.get(text_endpoint, headers=headers).json()
    return result


def transcribe_other_lang(file, token, lang):
    headers = {'authorization': token}
    url_response = requests.post('https://api.assemblyai.com/v2/upload',
                                 headers=headers,
                                 data=file)

    url = url_response.json()["upload_url"]

    id_endpoint = "https://api.assemblyai.com/v2/transcript"

    json = {
        "audio_url": url,
        "language_code": lang
    }

    headers = {"authorization": token, "content-type": "application/json"}
    id_response = requests.post(id_endpoint, json=json, headers=headers)
    transcribe_id = id_response.json()['id']
    text_endpoint = "https://api.assemblyai.com/v2/transcript/"+transcribe_id
    headers = {"authorization": token, }
    result = requests.get(text_endpoint, headers=headers).json()

    while result.get("status") != 'completed':
        if result.get("status") == 'error':
            return "error"
        text_endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {"authorization": token}
        result = requests.get(text_endpoint, headers=headers).json()
    return result

# ------------------------------------------


def json_data_extraction(result, fname):
    audindex = pd.json_normalize(result['words'])
    audindex['file_name'] = fname
    chapters = pd.json_normalize(result['chapters'])
    topics = pd.json_normalize(result['iab_categories_result']['results'])
    topics['label_1'] = topics['labels'].apply(lambda x: x[0]['label'])
    topics['label_2'] = topics['labels'].apply(
        lambda x: x[1]['label'] if len(x) > 1 else 'none')
    topics['label_3'] = topics['labels'].apply(
        lambda x: x[2]['label'] if len(x) > 2 else 'none')
    highlights = pd.json_normalize(result['auto_highlights_result']['results'])
    highlights = highlights.text.unique()
    audindex['summary'] = ""
    audindex['headline'] = ""
    audindex['gist'] = ""
    audindex['label_1'] = ""
    audindex['label_2'] = ""
    audindex['label_3'] = ""
    speakers = list(audindex.speaker)
    previous_speaker = 'A'
    l = len(speakers)
    i = 0
    speaker_seq_list = list()
    for index, new_speaker in enumerate(speakers):
        if index > 0:
            previous_speaker = speakers[index - 1]
        if new_speaker != previous_speaker:
            i += 1
        speaker_seq_list.append(i)
    audindex['sequence'] = speaker_seq_list

    for j in range(0, len(chapters)):
        for i in range(0, len(audindex)):
            if audindex.iloc[i]['start'] >= chapters.iloc[j]['start'] and audindex.iloc[i]['end'] <= chapters.iloc[j]['end']:
                audindex.loc[i, 'summary'] = chapters.iloc[j]['summary']
                audindex.loc[i, 'headline'] = chapters.iloc[j]['headline']
                audindex.loc[i, 'gist'] = chapters.iloc[j]['gist']

    for j in range(0, len(topics)):
        try:
            for i in range(0, len(audindex)):
                if audindex.iloc[i]['start'] >= topics.iloc[j]['timestamp.start'] and audindex.iloc[i]['end'] <= topics.iloc[j+1]['timestamp.start']:
                    audindex.loc[i, 'label_1'] = topics.iloc[j]['label_1']
                    audindex.loc[i, 'label_2'] = topics.iloc[j]['label_2']
                    audindex.loc[i, 'label_3'] = topics.iloc[j]['label_3']
        except:
            for i in range(0, len(audindex)):
                if audindex.iloc[i]['start'] >= topics.iloc[j]['timestamp.start']:
                    audindex.loc[i, 'label_1'] = topics.iloc[j]['label_1']
                    audindex.loc[i, 'label_2'] = topics.iloc[j]['label_2']
                    audindex.loc[i, 'label_3'] = topics.iloc[j]['label_3']

    group = ['file_name', 'speaker', 'summary', 'headline',
             'gist', 'label_1', 'label_2', 'label_3', 'sequence']
    df = pd.DataFrame(audindex.groupby(group).agg(
        utter=('text', ' '.join), start_time=('start', 'min'), end_time=('end', 'max')))
    df.reset_index(inplace=True)
    df['key_phrase'] = 'none'
    for x in highlights:
        df.loc[(df.utter.str.contains(x)), 'key_phrase'] = x
    df.sort_values(by=['start_time'], inplace=True)

    group_doc = ['speaker', 'sequence']
    df_doc = pd.DataFrame(audindex.groupby(group_doc).agg(
        utter=('text', ' '.join), start_time=('start', 'min')))
    df_doc.reset_index(inplace=True)
    df_doc.sort_values(by=['start_time'], inplace=True)
    df_doc_grouped = df_doc[['speaker', 'utter']]

    doc = docx.Document()
    t = doc.add_table(df_doc_grouped.shape[0]+1, df_doc_grouped.shape[1])
    for j in range(df_doc_grouped.shape[-1]):
        t.cell(0, j).text = df_doc_grouped.columns[j]
    for i in range(df_doc_grouped.shape[0]):
        for j in range(df_doc_grouped.shape[-1]):
            t.cell(i+1, j).text = str(df_doc_grouped.values[i, j])

    return df, doc


def start_transcribe(project_name, files, dir_name, use_id, username, lang):
    config = Connect.config()
    mydb = Connect.db()
    s3 = Connect.s3()
    tokens = eval(config['AssemblyAI']['tokens'])
    os.makedirs("documents/"+dir_name+"/"+project_name, exist_ok=True)
    html = ""
    exception_html = ""
    try:
        for file in files:
            fname = file.filename
            pre_cursor = mydb.cursor()
            pre_cursor.execute(
                "select username,project_name,fname,doc_download_link,csv_download_link,transcription_time from Records where project_name ='"+project_name+"' and fname ='"+fname+"';")
            table_rows = pre_cursor.fetchall()
            for re in table_rows:
                html += "<tr><td>"+re[0]+"</td><td>" + \
                    re[1]+"</td><td>"+re[2]+"</td>"
                html += "<td><a href='"+re[3] + \
                    "' class='dwn-link'>Download</a></td>"
                html += "<td><a href='"+re[4] + \
                    "' class='dwn-link'>Download</a></td>"
                html += "<td>"+re[5]+"</td><td>existing</td></tr>"
            file_list = []
            for row in table_rows:
                file_list.append(row[2])
                exception_html += row[2] + \
                    ":&nbsp;&nbsp;  already exists in the database<br>"
            if fname not in file_list:
                transcription_start_time = time.time()
                token = tokens[random.randint(0, len(tokens)-1)]
                if lang == "en":
                    result = transcribe_english(file.file, token)
                else:
                    result = transcribe_other_lang(file.file, token, lang)
                if result == "error":
                    html += "<tr><td>"+username+"</td><td>"+project_name+"</td><td>" + \
                        fname+"</td><td>-</td><td>-</td><td>-</td><td>error</td></tr>"
                    exception_html += fname+": please check the file and try again <br>"
                    error_log = open("logs/errors.log", "a+")
                    error_str = datetime.now().strftime("%m-%d-%Y, %I:%M %p")+":     "+username + \
                        ":     ["+fname + \
                        "]       : AssemblyAI Processing Error"
                    error_log.write(error_str+"\n")
                    error_log.close()
                    return HTMLResponse(content=html_completed(html, exception_html, username, use_id), status_code=200)
                transcription_time = str(
                    round((time.time()-transcription_start_time)/60, 2))+" mins"
                print("transcription completed")
                if lang == "en":
                    df, doc = json_data_extraction(result, fname)
                    df_path = "documents/"+dir_name+"/"+project_name+"/"+fname+".csv"
                    doc_path = "documents/"+dir_name+"/"+project_name+"/"+fname+".docx"
                    df.to_csv(df_path, index=False)
                    doc.save(doc_path)

                    s3.upload_file(doc_path,
                                   config['S3SETTINGS']['bucket'],
                                   config['S3SETTINGS']['folder']+project_name+"/"+fname+".docx")

                    s3.upload_file(df_path,
                                   config['S3SETTINGS']['bucket'],
                                   config['S3SETTINGS']['folder']+project_name+"/"+fname+".csv")

                    os.remove(doc_path)
                    os.remove(df_path)

                    endpoint = "http://ec2-54-164-248-248.compute-1.amazonaws.com/download/"
                    doc_download_link = endpoint+project_name+"/"+fname+".docx"
                    csv_download_link = endpoint+project_name+"/"+fname+".csv"

                    record_cursor = mydb.cursor()
                    query = "INSERT into Records(date_time,username,project_name,fname,doc_download_link,csv_download_link,transcription_time,language) values(%s,%s,%s,%s,%s,%s,%s,%s)"
                    record = [(datetime.now(), username, project_name, fname,
                               doc_download_link, csv_download_link, transcription_time, "English"),]
                    record_cursor.executemany(query, record)
                    mydb.commit()

                    html += "<tr><td>"+username+"</td><td>"+project_name+"</td><td>"+fname+"</td>"
                    html += "<td><a href='"+doc_download_link + \
                        "' class='dwn-link'>Download</a></td>"
                    html += "<td><a href='"+csv_download_link + \
                        "' class='dwn-link'>Download</a></td>"
                    html += "<td>"+transcription_time+"</td><td>success</td></tr>"
                else:
                    text = result['text']
                    text_list = text.split(".")
                    translated_text_list = []
                    for text in text_list:
                        translated = GoogleTranslator(
                            source='auto', target='en').translate(text)
                        translated_text_list.append(translated)
                    translated_text = ". ".join(translated_text_list)
                    doc = docx.Document()
                    doc.add_heading(project_name+": "+fname, level=1)
                    doc.add_paragraph(result['text'])
                    doc_path = "documents/"+dir_name+"/"+project_name+"/"+fname+".docx"
                    doc.save(doc_path)
                    doc1 = docx.Document()
                    doc1.add_heading(project_name+": "+fname, level=1)
                    doc1.add_paragraph(translated_text)
                    tdoc_path = "documents/"+dir_name+"/"+project_name+"/translated_"+fname+".docx"
                    doc1.save(tdoc_path)

                    s3.upload_file(doc_path,
                                   config['S3SETTINGS']['bucket'],
                                   config['S3SETTINGS']['folder']+project_name+"/"+fname+".docx")

                    s3.upload_file(tdoc_path,
                                   config['S3SETTINGS']['bucket'],
                                   config['S3SETTINGS']['folder']+project_name+"/translated_"+fname+".docx")

                    os.remove(doc_path)
                    os.remove(tdoc_path)

                    endpoint = "http://ec2-54-164-248-248.compute-1.amazonaws.com/download/"
                    doc_download_link = endpoint+project_name+"/"+fname+".docx"
                    trans_download_link = endpoint+project_name+"/translated_"+fname+".docx"

                    record_cursor = mydb.cursor()
                    query = "INSERT into Records(date_time,username,project_name,fname,doc_download_link,csv_download_link,transcription_time,language) values(%s,%s,%s,%s,%s,%s,%s,%s)"
                    record = [(datetime.now(), username, project_name, fname,
                               doc_download_link, trans_download_link, transcription_time, "Non English"),]
                    record_cursor.executemany(query, record)
                    mydb.commit()

                    html += "<tr><td>"+username+"</td><td>"+project_name+"</td><td>"+fname+"</td>"
                    html += "<td><a href='"+doc_download_link + \
                        "' class='dwn-link'>Download</a></td>"
                    html += "<td><a href='"+trans_download_link + \
                        "' class='dwn-link'>Download</a></td>"
                    html += "<td>"+transcription_time+"</td><td>success</td></tr>"
    except Exception as e:
        error_log = open("logs/error_log.txt", "a+")
        error_str = datetime.now().strftime("%m-%d-%Y, %I:%M %p") + \
            ": "+username+": ["+fname+"] "
        error_log.write(error_str+str(type(e))+"\n")
        error_log.write(str(e)+"\n")
        error_log.close()
        print(e)
        html += "<tr><td>"+username+"</td><td>"+project_name+"</td><td>" + \
            fname+"</td><td>-</td><td>-</td><td>-</td><td>error</td></tr>"
        exception_html += "error: "+str(e)
    shutil.rmtree("documents/"+dir_name)
    return html, exception_html, username, use_id
